# -*- coding: utf-8 -*-
"""
Qbix Centre — Complete Web Application
Runs on Railway. Serves both the public website and the management app.
Data stored in qbix_data.json (committed to Railway or on persistent volume).
"""

import json
import os
import secrets
import hashlib
import hmac
import smtplib
import ssl
import threading
import time
import pyotp
from datetime import datetime, timedelta
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from functools import wraps
from pathlib import Path

from flask import (Flask, render_template, request, jsonify, redirect,
                   url_for, session, flash, send_file, abort)
from itsdangerous import URLSafeTimedSerializer, BadSignature, SignatureExpired
from dotenv import load_dotenv

load_dotenv()

# ── App setup ────────────────────────────────────────────────────────────────
app = Flask(__name__, static_folder='static', static_url_path='/static')
app.secret_key = os.environ.get('SECRET_KEY', secrets.token_hex(32))
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=30)

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE_DIR  = Path(__file__).parent
DATA_FILE = BASE_DIR / os.environ.get('DATA_FILE', 'qbix_data.json')
BACKUP_DIR = BASE_DIR / 'backups'

# ── Config from environment ───────────────────────────────────────────────────
ADMIN_USERNAME    = os.environ.get('ADMIN_USERNAME', 'admin')
ADMIN_EMAIL       = os.environ.get('ADMIN_EMAIL', 'qbixcentre@outlook.com')
ADMIN_PHONE       = os.environ.get('ADMIN_PHONE', '4787379107')
APP_URL           = os.environ.get('APP_URL', 'http://localhost:5000')
FROM_EMAIL        = os.environ.get('FROM_EMAIL', 'noreply@qbixcentre.com')
FROM_NAME         = os.environ.get('FROM_NAME', 'Qbix Centre')
ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')
GA_MEASUREMENT_ID = os.environ.get('GA_MEASUREMENT_ID', '')

# TOTP secret for 2FA (generated once and stored in env)
TOTP_SECRET = os.environ.get('TOTP_SECRET', pyotp.random_base32())

# Serializer for signed tokens (onboarding links, booking tokens)
serializer = URLSafeTimedSerializer(app.secret_key)

# In-memory stores (fine for single-instance Railway deployment)
_pending_2fa   = {}   # session_id -> {code, expires, purpose}
_booking_tokens = {}  # token -> member_email
_onboard_tokens = {}  # token -> {name, email, expires}

# ── Default data ──────────────────────────────────────────────────────────────
DEFAULT_DATA = {
    "offices": [
        {"id":"o1","num":"11","status":"Occupied","member":"HighBar Accounting","tenantStart":"9/1/2025","sqft":103,"dormer":None,"listDues":None},
        {"id":"o2","num":"12","status":"Occupied","member":"Pinnacle Accounting","tenantStart":"9/1/2025","sqft":139,"dormer":None,"listDues":None},
        {"id":"o3","num":"13","status":"Occupied","member":"Retail 1","tenantStart":"12/11/2018","sqft":196,"dormer":None,"listDues":None},
        {"id":"o4","num":"14","status":"Occupied","member":"Pinnacle Accounting","tenantStart":"9/1/2025","sqft":147,"dormer":None,"listDues":None},
        {"id":"o5","num":"15","status":"Occupied","member":"Pinnacle Accounting","tenantStart":"9/1/2025","sqft":150,"dormer":None,"listDues":None},
        {"id":"o6","num":"16","status":"Occupied","member":"Biren Patel Engineering","tenantStart":"6/8/2019","sqft":245,"dormer":None,"listDues":None},
        {"id":"o7","num":"17","status":"Occupied","member":"Biren Patel Engineering","tenantStart":"6/8/2019","sqft":196,"dormer":None,"listDues":None},
        {"id":"o8","num":"18","status":"Occupied","member":"Pettis Group","tenantStart":"12/1/2022","sqft":209,"dormer":None,"listDues":None},
        {"id":"o9","num":"19","status":"Occupied","member":"HighBar Accounting","tenantStart":"9/1/2025","sqft":176,"dormer":None,"listDues":None},
        {"id":"o10","num":"19A","status":"Occupied","member":"Preferred Provider Network","tenantStart":"9/2/2025","sqft":140,"dormer":None,"listDues":None},
        {"id":"o11","num":"19B","status":"Vacant","member":"","tenantStart":"","sqft":140,"dormer":None,"listDues":725},
        {"id":"o12","num":"21","status":"Occupied","member":"Gilbert Gomez CPA","tenantStart":"9/2/2025","sqft":90,"dormer":31,"listDues":None},
        {"id":"o13","num":"22","status":"Occupied","member":"McLendon Law","tenantStart":"10/1/2021","sqft":90,"dormer":31,"listDues":None},
        {"id":"o14","num":"23","status":"Occupied","member":"NAG Enterprise Group, LLC","tenantStart":"6/1/2025","sqft":99,"dormer":None,"listDues":None},
        {"id":"o15","num":"24","status":"Occupied","member":"HTNB Corp","tenantStart":"5/10/2025","sqft":90,"dormer":31,"listDues":None},
        {"id":"o16","num":"25","status":"Occupied","member":"McLendon Law","tenantStart":"10/1/2021","sqft":87,"dormer":None,"listDues":None},
        {"id":"o17","num":"26","status":"Vacant","member":"","tenantStart":"","sqft":87,"dormer":None,"listDues":500},
        {"id":"o18","num":"27","status":"Occupied","member":"Care Forth","tenantStart":"5/1/2021","sqft":87,"dormer":None,"listDues":None},
        {"id":"o19","num":"28","status":"Occupied","member":"Joshua David Nicholson","tenantStart":"10/1/2022","sqft":87,"dormer":None,"listDues":None},
        {"id":"o20","num":"29A","status":"Occupied","member":"Larry Fouche","tenantStart":"10/1/2025","sqft":142,"dormer":None,"listDues":None},
        {"id":"o21","num":"29B","status":"Vacant","member":"","tenantStart":"","sqft":158,"dormer":None,"listDues":725},
        {"id":"o22","num":"31","status":"Occupied","member":"Preferred Provider Network","tenantStart":"9/2/2025","sqft":139,"dormer":None,"listDues":None},
        {"id":"o23","num":"32","status":"Occupied","member":"National Youth Advocate Program","tenantStart":"5/1/2025","sqft":144,"dormer":None,"listDues":None},
        {"id":"o24","num":"33","status":"Occupied","member":"Wilson PC","tenantStart":"7/1/2023","sqft":161,"dormer":None,"listDues":None},
        {"id":"o25","num":"34","status":"Occupied","member":"Ram Bay","tenantStart":"10/1/2024","sqft":189,"dormer":None,"listDues":None},
        {"id":"o26","num":"35","status":"Occupied","member":"Rid A Critter","tenantStart":"2/1/2026","sqft":81,"dormer":None,"listDues":None},
        {"id":"o27","num":"36","status":"Occupied","member":"Ram Bay","tenantStart":"10/1/2024","sqft":136,"dormer":None,"listDues":None},
    ],
    "members": [],
    "occupants": [],
    "waitlist": [],
    "bookings": [],
    "templates": [
        {"id":"t1","name":"Power Outage","subject":"Power Outage Notice — Qbix Centre","body":"Dear {name},\n\nPlease be advised that Qbix Centre is currently experiencing a power outage. We are working to restore power as quickly as possible.\n\nWe apologize for any inconvenience and will keep you updated.\n\nThank you for your patience,\nQbix Centre Management"},
        {"id":"t2","name":"Monthly Dues Reminder","subject":"Monthly Dues Reminder — Qbix Centre","body":"Dear {name},\n\nThis is a friendly reminder that your monthly dues of {dues} are due. Please arrange payment at your earliest convenience.\n\nThank you,\nQbix Centre Management"},
        {"id":"t3","name":"Building Maintenance","subject":"Planned Maintenance Notice — Qbix Centre","body":"Dear {name},\n\nWe wanted to notify you that Qbix Centre will be undergoing scheduled maintenance. During this time, some services may be temporarily unavailable.\n\nBest regards,\nQbix Centre Management"},
        {"id":"t4","name":"General Notice","subject":"Important Notice from Qbix Centre","body":"Dear {name},\n\n[Your message here]\n\nBest regards,\nQbix Centre Management"},
    ],
    "lastBackup": "",
    "newsletter": [],
}


# ── Data helpers ──────────────────────────────────────────────────────────────
_data_lock = threading.Lock()

def load_data():
    if DATA_FILE.exists():
        with open(DATA_FILE, 'r', encoding='utf-8') as f:
            d = json.load(f)
        # Migrate: ensure new fields exist
        d.setdefault('bookings', [])
        d.setdefault('newsletter', [])
        for m in d.get('members', []):
            m.setdefault('attachments', [])
            m.setdefault('discount', 0)
            m.setdefault('agreementSent', '')
            m.setdefault('agreementSigned', '')
        for p in d.get('occupants', []):
            p.setdefault('dlAttachment', None)
        return d
    return json.loads(json.dumps(DEFAULT_DATA))

def save_data(data):
    with _data_lock:
        with open(DATA_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)

def get_db():
    return load_data()

def net_dues(member):
    return max(0, (member.get('dues') or 0) - (member.get('discount') or 0))

def offices_for(data, name):
    return [o['num'] for o in data['offices'] if o.get('member') == name]

def hours_included(data, member_name):
    """Members get 6 hours per office they hold."""
    return len(offices_for(data, member_name)) * 6


# ── Email helper ──────────────────────────────────────────────────────────────
def send_email(to_email, to_name, subject, html_body, text_body=None):
    """Send email via SMTP. Configure SMTP_* env vars."""
    smtp_host = os.environ.get('SMTP_HOST', 'smtp.gmail.com')
    smtp_port = int(os.environ.get('SMTP_PORT', 587))
    smtp_user = os.environ.get('SMTP_USER', '')
    smtp_pass = os.environ.get('SMTP_PASS', '')

    if not smtp_user:
        print(f"[EMAIL] Would send to {to_email}: {subject}")
        return True

    msg = MIMEMultipart('alternative')
    msg['Subject'] = subject
    msg['From']    = f"{FROM_NAME} <{FROM_EMAIL}>"
    msg['To']      = f"{to_name} <{to_email}>"

    if text_body:
        msg.attach(MIMEText(text_body, 'plain'))
    msg.attach(MIMEText(html_body, 'html'))

    try:
        context = ssl.create_default_context()
        with smtplib.SMTP(smtp_host, smtp_port, timeout=10) as server:
            server.starttls(context=context)
            server.login(smtp_user, smtp_pass)
            server.sendmail(FROM_EMAIL, to_email, msg.as_string())
        return True
    except Exception as e:
        print(f"[EMAIL ERROR] {e}")
        return False


def send_sms_code(phone, code):
    """Send SMS via email-to-SMS gateway (free) or Twilio if configured."""
    # Email-to-SMS gateways (free, carrier dependent)
    gateways = {
        'att':      f'{phone}@txt.att.net',
        'verizon':  f'{phone}@vtext.com',
        'tmobile':  f'{phone}@tmomail.net',
        'sprint':   f'{phone}@messaging.sprintpcs.com',
    }
    # Default: try AT&T gateway (most common in GA)
    carrier = os.environ.get('ADMIN_CARRIER', 'att')
    sms_email = gateways.get(carrier, gateways['att'])

    msg = MIMEText(f"Qbix Centre login code: {code}")
    msg['Subject'] = ''
    msg['From']    = FROM_EMAIL
    msg['To']      = sms_email

    smtp_host = os.environ.get('SMTP_HOST', 'smtp.gmail.com')
    smtp_port = int(os.environ.get('SMTP_PORT', 587))
    smtp_user = os.environ.get('SMTP_USER', '')
    smtp_pass = os.environ.get('SMTP_PASS', '')

    if not smtp_user:
        print(f"[SMS] Would send code {code} to {phone}")
        return True

    try:
        context = ssl.create_default_context()
        with smtplib.SMTP(smtp_host, smtp_port) as server:
            server.starttls(context=context)
            server.login(smtp_user, smtp_pass)
            server.sendmail(FROM_EMAIL, sms_email, msg.as_string())
        return True
    except Exception as e:
        print(f"[SMS ERROR] {e}")
        return False


# ── Auth helpers ──────────────────────────────────────────────────────────────
def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def check_password(password):
    stored = os.environ.get('ADMIN_PASSWORD_HASH', '')
    if not stored:
        # First run — accept any password and prompt setup
        return True
    return hmac.compare_digest(hash_password(password), stored)

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('admin_authenticated'):
            return redirect(url_for('admin_login'))
        return f(*args, **kwargs)
    return decorated

def generate_code():
    return str(secrets.randbelow(900000) + 100000)  # 6-digit


# ── Template context ─────────────────────────────────────────────────────────
@app.context_processor
def inject_globals():
    return {'now': datetime.now()}


# ── Health check ──────────────────────────────────────────────────────────────
@app.route('/health')
def health():
    return jsonify({'ok': True, 'time': datetime.now().isoformat()})


# ══════════════════════════════════════════════════════════════════════════════
# PUBLIC WEBSITE ROUTES
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/')
def home():
    data = get_db()
    occupied = len([o for o in data['offices'] if o['status'] == 'Occupied'])
    total    = len(data['offices'])
    vacant   = [o for o in data['offices'] if o['status'] == 'Vacant']
    return render_template('public/home.html',
        occupied=occupied, total=total, vacant=vacant,
        ga_id=GA_MEASUREMENT_ID)

@app.route('/offices')
def offices_page():
    data = get_db()
    vacant = [o for o in data['offices'] if o['status'] == 'Vacant']
    return render_template('public/offices.html', vacant=vacant, ga_id=GA_MEASUREMENT_ID)

@app.route('/amenities')
def amenities():
    return render_template('public/amenities.html', ga_id=GA_MEASUREMENT_ID)

@app.route('/contact')
def contact():
    return render_template('public/contact.html', ga_id=GA_MEASUREMENT_ID)

@app.route('/contact', methods=['POST'])
def contact_submit():
    name    = request.form.get('name', '')
    email   = request.form.get('email', '')
    phone   = request.form.get('phone', '')
    message = request.form.get('message', '')
    # Email to admin
    send_email(
        ADMIN_EMAIL, 'Qbix Centre Admin',
        f'Website enquiry from {name}',
        f'<p><b>Name:</b> {name}<br><b>Email:</b> {email}<br><b>Phone:</b> {phone}</p><p>{message}</p>',
        f'Name: {name}\nEmail: {email}\nPhone: {phone}\n\n{message}'
    )
    flash('Thank you! We will be in touch shortly.', 'success')
    return redirect(url_for('contact'))

@app.route('/news')
def news():
    data = get_db()
    posts = sorted(data.get('newsletter', []), key=lambda x: x.get('date',''), reverse=True)
    return render_template('public/news.html', posts=posts, ga_id=GA_MEASUREMENT_ID)

@app.route('/news/<post_id>')
def news_post(post_id):
    data = get_db()
    post = next((p for p in data.get('newsletter', []) if p['id'] == post_id), None)
    if not post:
        abort(404)
    return render_template('public/news_post.html', post=post, ga_id=GA_MEASUREMENT_ID)


# ══════════════════════════════════════════════════════════════════════════════
# ONBOARDING FLOW (public — for prospective members)
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/onboard')
def onboard_home():
    """Landing page linked from website — general interest form."""
    return render_template('public/onboard_home.html', ga_id=GA_MEASUREMENT_ID)

@app.route('/onboard/<token>')
def onboard(token):
    """Personalized onboarding link sent by admin."""
    if token not in _onboard_tokens:
        return render_template('public/onboard_expired.html')
    info = _onboard_tokens[token]
    if datetime.now() > info['expires']:
        del _onboard_tokens[token]
        return render_template('public/onboard_expired.html')
    return render_template('public/onboard.html', token=token, info=info, ga_id=GA_MEASUREMENT_ID)

@app.route('/onboard/<token>/submit', methods=['POST'])
def onboard_submit(token):
    if token not in _onboard_tokens:
        return jsonify({'ok': False, 'error': 'Link expired'}), 400

    data = get_db()
    form = request.form

    # Create pending member
    member_id = '_' + secrets.token_hex(4)
    member = {
        'id':        member_id,
        'name':      form.get('company') or form.get('firstName') + ' ' + form.get('lastName'),
        'status':    'Pending',
        'start':     form.get('startDate', ''),
        'end':       '',
        'dues':      0,
        'discount':  0,
        'deposit':   0,
        'notes':     form.get('notes', ''),
        'email':     form.get('email', ''),
        'phone':     form.get('phone', ''),
        'address':   form.get('address', ''),
        'city':      form.get('city', ''),
        'state':     form.get('state', ''),
        'zip':       form.get('zip', ''),
        'website':   form.get('website', ''),
        'attachments':      [],
        'agreementSent':    '',
        'agreementSigned':  '',
        'emergencyName':    form.get('emergencyName', ''),
        'emergencyPhone':   form.get('emergencyPhone', ''),
        'emergencyRel':     form.get('emergencyRel', ''),
        'onboardedAt':      datetime.now().isoformat(),
    }
    data['members'].append(member)

    # Create pending occupant
    occ_id = '_' + secrets.token_hex(4)
    occupant = {
        'id':          occ_id,
        'name':        form.get('firstName', '') + ' ' + form.get('lastName', ''),
        'company':     member['name'],
        'phone':       form.get('phone', ''),
        'email':       form.get('email', ''),
        'office':      '',
        'endDate':     '',
        'status':      'Pending',
        'dlAttachment': None,
    }
    data['occupants'].append(occupant)
    save_data(data)

    # Notify admin
    send_email(
        ADMIN_EMAIL, 'Qbix Centre Admin',
        f'New onboarding submission: {member["name"]}',
        f'<p>A new prospect has completed the onboarding form.</p>'
        f'<p><b>Name:</b> {occupant["name"]}<br>'
        f'<b>Company:</b> {member["name"]}<br>'
        f'<b>Email:</b> {member["email"]}<br>'
        f'<b>Phone:</b> {member["phone"]}</p>'
        f'<p>Log in to your Qbix Centre dashboard to review and activate.</p>'
    )

    del _onboard_tokens[token]
    return jsonify({'ok': True})


# ══════════════════════════════════════════════════════════════════════════════
# CONFERENCE ROOM BOOKING (member-facing)
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/book')
def book_home():
    return render_template('public/book_home.html', ga_id=GA_MEASUREMENT_ID)

@app.route('/book/request-code', methods=['POST'])
def book_request_code():
    email = request.json.get('email', '').strip().lower()
    data  = get_db()

    # Check if active member with this email
    member = next((m for m in data['members']
                   if m.get('email','').lower() == email
                   and m.get('status') == 'Active'), None)

    # Also check occupants - find their member company
    if not member:
        occ = next((o for o in data['occupants']
                    if o.get('email','').lower() == email
                    and o.get('status') == 'Active'), None)
        if occ:
            member = next((m for m in data['members']
                          if m.get('name') == occ.get('company')
                          and m.get('status') == 'Active'), None)
            if not member:
                # Create a pseudo-member from occupant
                member = {'name': occ.get('name'), 'email': email, 'status': 'Active'}

    if not member:
        return jsonify({'ok': False, 'error': 'Email not found in our active member list.'})

    code = generate_code()
    token = secrets.token_urlsafe(32)
    _pending_2fa[token] = {
        'code':    code,
        'email':   email,
        'name':    member['name'],
        'expires': datetime.now() + timedelta(minutes=10),
    }

    send_email(
        email, member['name'],
        'Your Qbix Centre booking code',
        f'<p>Your conference room booking code is:</p>'
        f'<h1 style="letter-spacing:8px;color:#2563eb">{code}</h1>'
        f'<p>This code expires in 10 minutes.</p>',
        f'Your Qbix Centre booking code: {code}\nExpires in 10 minutes.'
    )

    return jsonify({'ok': True, 'token': token})

@app.route('/book/verify', methods=['POST'])
def book_verify():
    token = request.json.get('token', '')
    code  = request.json.get('code', '').strip()

    entry = _pending_2fa.get(token)
    if not entry:
        return jsonify({'ok': False, 'error': 'Invalid or expired session.'})
    if datetime.now() > entry['expires']:
        del _pending_2fa[token]
        return jsonify({'ok': False, 'error': 'Code expired. Please request a new one.'})
    if entry['code'] != code:
        return jsonify({'ok': False, 'error': 'Incorrect code.'})

    # Issue booking session token
    booking_token = secrets.token_urlsafe(32)
    _booking_tokens[booking_token] = {
        'email': entry['email'],
        'name':  entry['name'],
        'expires': datetime.now() + timedelta(hours=2),
    }
    del _pending_2fa[token]
    return jsonify({'ok': True, 'bookingToken': booking_token})

@app.route('/book/calendar')
def book_calendar():
    bt = request.args.get('token', '')
    entry = _booking_tokens.get(bt)
    if not entry or datetime.now() > entry['expires']:
        return redirect(url_for('book_home'))

    data = get_db()
    member = next((m for m in data['members']
                   if m.get('email','').lower() == entry['email'].lower()), None)
    if not member:
        return redirect(url_for('book_home'))

    included = hours_included(data, member['name'])
    return render_template('public/book_calendar.html',
        token=bt,
        member_name=entry['name'],
        included_hours=included,
        ga_id=GA_MEASUREMENT_ID)

@app.route('/book/slots')
def book_slots():
    """Return booked slots for a given month."""
    bt = request.args.get('token', '')
    if bt not in _booking_tokens:
        return jsonify({'ok': False}), 401

    year  = int(request.args.get('year',  datetime.now().year))
    month = int(request.args.get('month', datetime.now().month))
    data  = get_db()

    slots = [b for b in data.get('bookings', [])
             if b.get('year') == year and b.get('month') == month
             and b.get('status') != 'cancelled']
    return jsonify({'ok': True, 'slots': slots})

@app.route('/book/create', methods=['POST'])
def book_create():
    bt = request.json.get('token', '')
    entry = _booking_tokens.get(bt)
    if not entry or datetime.now() > entry['expires']:
        return jsonify({'ok': False, 'error': 'Session expired'}), 401

    data   = get_db()
    member = next((m for m in data['members']
                   if m.get('email','').lower() == entry['email'].lower()), None)
    if not member:
        return jsonify({'ok': False, 'error': 'Member not found'}), 400

    date_str   = request.json.get('date', '')    # YYYY-MM-DD
    start_time = request.json.get('start', '')   # HH:MM
    end_time   = request.json.get('end', '')     # HH:MM
    title      = request.json.get('title', 'Meeting')

    try:
        dt = datetime.strptime(date_str, '%Y-%m-%d')
    except ValueError:
        return jsonify({'ok': False, 'error': 'Invalid date'}), 400

    # Check for conflicts
    for b in data.get('bookings', []):
        if (b.get('date') == date_str
                and b.get('status') != 'cancelled'
                and not (end_time <= b.get('start','') or start_time >= b.get('end',''))):
            return jsonify({'ok': False, 'error': 'That time slot is already booked.'})

    booking_id = '_' + secrets.token_hex(4)
    booking = {
        'id':          booking_id,
        'memberName':  member['name'],
        'memberEmail': entry['email'],
        'date':        date_str,
        'year':        dt.year,
        'month':       dt.month,
        'start':       start_time,
        'end':         end_time,
        'title':       title,
        'status':      'confirmed',
        'createdAt':   datetime.now().isoformat(),
    }
    data.setdefault('bookings', []).append(booking)
    save_data(data)

    # Confirmation email to member
    send_email(
        entry['email'], entry['name'],
        f'Conference Room Booking Confirmed — {date_str}',
        f'<p>Hi {entry["name"]},</p>'
        f'<p>Your conference room booking is confirmed:</p>'
        f'<ul><li><b>Date:</b> {date_str}</li>'
        f'<li><b>Time:</b> {start_time} – {end_time}</li>'
        f'<li><b>Room:</b> Qbix Centre Conference Room</li></ul>'
        f'<p>We look forward to seeing you!</p>'
        f'<p style="color:#666;font-size:12px">500A Northside Crossing, Macon, GA 31210</p>',
    )

    # Schedule reminder (24h before) in background thread
    def send_reminder():
        try:
            booking_dt = datetime.strptime(f'{date_str} {start_time}', '%Y-%m-%d %H:%M')
            reminder_dt = booking_dt - timedelta(hours=24)
            wait = (reminder_dt - datetime.now()).total_seconds()
            if wait > 0:
                time.sleep(wait)
            send_email(
                entry['email'], entry['name'],
                f'Reminder: Conference Room Tomorrow at {start_time}',
                f'<p>Hi {entry["name"]},</p>'
                f'<p>Just a reminder that you have the Qbix Centre conference room booked tomorrow:</p>'
                f'<ul><li><b>Date:</b> {date_str}</li>'
                f'<li><b>Time:</b> {start_time} – {end_time}</li></ul>'
                f'<p>See you tomorrow!</p>',
            )
        except Exception as e:
            print(f'Reminder error: {e}')

    threading.Thread(target=send_reminder, daemon=True).start()

    return jsonify({'ok': True, 'booking': booking})

@app.route('/book/cancel', methods=['POST'])
def book_cancel():
    bt = request.json.get('token', '')
    entry = _booking_tokens.get(bt)
    if not entry or datetime.now() > entry['expires']:
        return jsonify({'ok': False}), 401

    booking_id = request.json.get('bookingId', '')
    data = get_db()
    booking = next((b for b in data.get('bookings', [])
                    if b['id'] == booking_id
                    and b['memberEmail'].lower() == entry['email'].lower()), None)
    if not booking:
        return jsonify({'ok': False, 'error': 'Booking not found'}), 404

    booking['status'] = 'cancelled'
    save_data(data)
    return jsonify({'ok': True})


# ══════════════════════════════════════════════════════════════════════════════
# ADMIN LOGIN (2-factor)
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    # Check if already fully authenticated
    if session.get('admin_authenticated'):
        return redirect(url_for('admin_dashboard'))

    if request.method == 'POST':
        username = request.form.get('username', '')
        password = request.form.get('password', '')

        if username == ADMIN_USERNAME and check_password(password):
            # Send 2FA code via SMS
            code  = generate_code()
            sid   = secrets.token_urlsafe(16)
            _pending_2fa[sid] = {
                'code':    code,
                'expires': datetime.now() + timedelta(minutes=10),
                'purpose': 'admin',
            }
            session['admin_2fa_sid'] = sid

            # Try SMS first, fall back to email
            sms_ok = send_sms_code(ADMIN_PHONE, code)
            if not sms_ok:
                # Send via email as fallback
                send_email(
                    ADMIN_EMAIL, 'Admin',
                    'Your Qbix Centre login code',
                    f'<p>Your login code is: <strong style="font-size:24px;letter-spacing:4px">{code}</strong></p><p>Expires in 10 minutes.</p>',
                    f'Your Qbix Centre login code: {code}\nExpires in 10 minutes.'
                )

            return redirect(url_for('admin_2fa'))
        else:
            flash('Invalid username or password.', 'error')

    return render_template('admin/login.html')

@app.route('/admin/2fa', methods=['GET', 'POST'])
def admin_2fa():
    sid = session.get('admin_2fa_sid')
    if not sid or sid not in _pending_2fa:
        return redirect(url_for('admin_login'))

    if request.method == 'POST':
        code  = request.form.get('code', '').strip()
        entry = _pending_2fa.get(sid)

        if not entry or datetime.now() > entry['expires']:
            flash('Code expired. Please log in again.', 'error')
            return redirect(url_for('admin_login'))

        if entry['code'] != code:
            flash('Incorrect code. Please try again.', 'error')
            return render_template('admin/2fa.html')

        # Success — fully authenticated
        del _pending_2fa[sid]
        session.pop('admin_2fa_sid', None)
        session['admin_authenticated'] = True
        session['admin_login_time']    = datetime.now().isoformat()
        session.permanent = True

        return redirect(url_for('admin_dashboard'))

    return render_template('admin/2fa.html')

@app.route('/admin/logout')
def admin_logout():
    session.clear()
    return redirect(url_for('home'))


# ══════════════════════════════════════════════════════════════════════════════
# ADMIN — MANAGEMENT APP
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/admin')
@login_required
def admin_dashboard():
    data = get_db()
    active  = [m for m in data['members'] if m['status'] == 'Active']
    pending = [m for m in data['members'] if m['status'] == 'Pending']
    occ     = len([o for o in data['offices'] if o['status'] == 'Occupied'])
    vac     = len([o for o in data['offices'] if o['status'] == 'Vacant'])
    gross   = sum((m.get('dues') or 0) for m in active)
    net_rev = sum(net_dues(m) for m in active)
    dep     = sum((m.get('deposit') or 0) for m in active)
    endings = sorted([m for m in data['members'] if m.get('end')],
                     key=lambda m: m['end'])[:5]
    vacant  = [o for o in data['offices'] if o['status'] == 'Vacant']

    # Upcoming bookings (next 7 days)
    today = datetime.now().date()
    upcoming = sorted(
        [b for b in data.get('bookings', [])
         if b.get('status') == 'confirmed'
         and datetime.strptime(b['date'], '%Y-%m-%d').date() >= today],
        key=lambda b: (b['date'], b['start'])
    )[:10]

    return render_template('admin/dashboard.html',
        active=active, pending=pending,
        occ=occ, vac=vac, gross=gross, net_rev=net_rev, dep=dep,
        endings=endings, vacant=vacant, upcoming_bookings=upcoming,
        total=len(data['offices']),
        data=data)

# ── Data API (used by admin JS frontend) ─────────────────────────────────────

@app.route('/admin/api/data')
@login_required
def api_data():
    return jsonify(get_db())

@app.route('/admin/api/save', methods=['POST'])
@login_required
def api_save():
    try:
        data = request.json
        save_data(data)
        return jsonify({'ok': True})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/admin/api/backup')
@login_required
def api_backup():
    try:
        BACKUP_DIR.mkdir(exist_ok=True)
        today = datetime.now().strftime('%Y-%m-%d')
        dest  = BACKUP_DIR / f'qbix-backup-{today}.json'
        import shutil
        shutil.copy2(DATA_FILE, dest)
        data = get_db()
        data['lastBackup'] = today
        save_data(data)
        return jsonify({'ok': True, 'path': str(dest)})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/admin/api/test-email')
@login_required
def test_email():
    """Send a test email to the admin address to verify SMTP works."""
    import os
    cfg = {
        'SMTP_HOST': os.environ.get('SMTP_HOST',''),
        'SMTP_PORT': os.environ.get('SMTP_PORT',''),
        'SMTP_USER': os.environ.get('SMTP_USER',''),
        'SMTP_PASS': '***' if os.environ.get('SMTP_PASS') else '(not set)',
        'FROM_EMAIL': os.environ.get('FROM_EMAIL',''),
        'ADMIN_EMAIL': os.environ.get('ADMIN_EMAIL',''),
        'ADMIN_PHONE': os.environ.get('ADMIN_PHONE',''),
        'ADMIN_CARRIER': os.environ.get('ADMIN_CARRIER',''),
    }
    ok = send_email(
        ADMIN_EMAIL, 'Rocky',
        'Qbix Centre — SMTP Test',
        '<h2>SMTP is working!</h2><p>Your Qbix Centre email configuration is correct.</p>',
        'SMTP is working! Your Qbix Centre email configuration is correct.'
    )
    return jsonify({'ok': ok, 'config': cfg,
                    'message': 'Email sent!' if ok else 'Email FAILED — check config'})



@app.route('/admin/api/onboard-link', methods=['POST'])
@login_required
def generate_onboard_link():
    name  = request.json.get('name', '')
    email = request.json.get('email', '')
    token = secrets.token_urlsafe(16)
    _onboard_tokens[token] = {
        'name':    name,
        'email':   email,
        'expires': datetime.now() + timedelta(days=7),
    }
    link = f'{APP_URL}/onboard/{token}'

    # Email the link to the prospect
    if email:
        send_email(
            email, name,
            'Welcome to Qbix Centre — Complete Your Application',
            f'<p>Hi {name},</p>'
            f'<p>Thank you for your interest in Qbix Centre! Please click the link below to complete your membership application:</p>'
            f'<p><a href="{link}" style="background:#2563eb;color:#fff;padding:12px 24px;border-radius:6px;text-decoration:none;display:inline-block">Complete My Application</a></p>'
            f'<p>This link expires in 7 days.</p>'
            f'<p>If you have any questions, reply to this email or call (478) 216-2876.</p>'
            f'<p>We look forward to welcoming you!</p>'
            f'<p>— The Qbix Centre Team</p>',
        )

    return jsonify({'ok': True, 'link': link, 'token': token})

# ── Agreement generator ───────────────────────────────────────────────────────

@app.route('/admin/api/generate-agreement/<member_id>')
@login_required
def generate_agreement(member_id):
    """Generate a complete, professional filled-in membership agreement as .docx"""
    data   = get_db()
    member = next((m for m in data['members'] if m['id'] == member_id), None)
    if not member:
        abort(404)

    offices     = offices_for(data, member['name'])
    office_str  = ', '.join(offices) if offices else 'TBD'
    dues_amt    = member.get('dues', 0) or 0
    deposit_amt = member.get('deposit', 0) or 0
    dues_str    = f"${dues_amt:,}/month"
    deposit_str = f"${deposit_amt:,}"
    start_str   = member.get('start', '') or '_______________'
    today_str   = datetime.now().strftime('%B %d, %Y')
    conf_hours  = len(offices) * 6

    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import io

        doc = DocxDocument()

        # Page margins
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1.25)
            section.right_margin  = Inches(1.25)

        def add_heading(text, level=1):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(11) if level == 1 else Pt(10)
            run.font.color.rgb = RGBColor(0x1a, 0x27, 0x44)
            # Bottom border
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '1a2744')
            pBdr.append(bottom)
            pPr.append(pBdr)
            return p

        def add_body(text):
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.size = Pt(10)
            return p

        def add_bullet(text):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(text)
            run.font.size = Pt(10)
            return p

        def add_field_row(label, value, bold_value=False):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            r1 = p.add_run(f'{label}:  ')
            r1.bold = True
            r1.font.size = Pt(10)
            r2 = p.add_run(value or '___________________________')
            r2.bold = bold_value
            r2.font.size = Pt(10)
            if bold_value:
                r2.font.color.rgb = RGBColor(0x1a, 0x27, 0x44)
            return p

        def add_sig_line(label, prefill=''):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            r1 = p.add_run(f'{label}:  ')
            r1.bold = True
            r1.font.size = Pt(10)
            r2 = p.add_run(prefill if prefill else '_' * 45)
            r2.font.size = Pt(10)
            return p

        # ── TITLE ──────────────────────────────────────────────────────────
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(2)
        tr = title.add_run('QBIX CENTRE')
        tr.bold = True
        tr.font.size = Pt(20)
        tr.font.color.rgb = RGBColor(0x1a, 0x27, 0x44)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.paragraph_format.space_after = Pt(2)
        sr = sub.add_run('Membership Agreement & House Guidelines')
        sr.font.size = Pt(12)
        sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        addr = doc.add_paragraph()
        addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        addr.paragraph_format.space_after = Pt(2)
        ar = addr.add_run('500A Northside Crossing, Macon, GA 31210  |  (478) 787-0532  |  qbixcentre.com')
        ar.font.size = Pt(9)
        ar.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        ver = doc.add_paragraph()
        ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ver.paragraph_format.space_after = Pt(16)
        vr = ver.add_run(f'Agreement Date: {today_str}')
        vr.font.size = Pt(9)
        vr.italic = True

        # ── MEMBER SUMMARY BOX ─────────────────────────────────────────────
        add_heading('Member Summary')
        add_field_row('Member / Company', member.get('name', ''), bold_value=True)
        add_field_row('Office(s) Assigned', office_str, bold_value=True)
        add_field_row('Monthly Dues', dues_str, bold_value=True)
        add_field_row('Refundable Deposit', deposit_str)
        add_field_row('Initial Term Start Date', start_str)
        add_field_row('Conference Room Hours Included', f'{conf_hours} hours/month')
        add_field_row('Contact Email', member.get('email', ''))
        add_field_row('Contact Phone', member.get('phone', ''))
        doc.add_paragraph()

        # ── SECTION 1 ──────────────────────────────────────────────────────
        add_heading('1. Membership & Fees')
        add_bullet(f'Monthly dues of {dues_str} are due on the 1st of each month via auto-draft through Bill.com. Payments received after the 5th are considered late.')
        add_bullet(f'A refundable deposit of {deposit_str} (equivalent to one month\'s dues) is required prior to move-in and will be returned less normal wear and tear and cost of unreturned keys.')
        add_bullet('A one-time setup fee of $200 is due at signing.')
        add_bullet(f'The initial term is six (6) months beginning {start_str}, automatically renewing on a month-to-month basis at the then-current rate unless terminated in accordance with Section 6.')
        add_bullet('Address Memberships are month-to-month from inception.')
        add_bullet('A non-refundable background check fee of $35 per cardholder is required prior to access being granted.')
        add_bullet('Additional key/fob holders: $150/month plus $35 background check fee. Both keyholders must be from the same company.')

        # ── SECTION 2 ──────────────────────────────────────────────────────
        add_heading('2. Access & Use')
        add_bullet('Members receive 24/7 access via card key/fob and a personal access code. Access codes are strictly confidential and must not be shared.')
        add_bullet(f'The conference room may be reserved online at qbixcentre.com at least 24 hours in advance. Each membership includes {conf_hours} hours per month at no charge; additional time is billed at $50/hour.')
        add_bullet('The workspace is for lawful, professional business purposes only. Sleeping, cooking meals, or conducting personal activities on the premises is not permitted.')
        add_bullet('Members are responsible for safeguarding their own confidential information and must respect the privacy and confidentiality of fellow members.')

        # ── SECTION 3 ──────────────────────────────────────────────────────
        add_heading('3. Amenities & Overage Charges')
        add_bullet('Included with all memberships: High-speed Wi-Fi/Ethernet (AT&T Gigabit Fiber), furnished workstations with sit/stand desks, kitchenette with Starbucks coffee and beverages, full-color laser printer/scanner, conference room, free parking, and janitorial service.')
        add_bullet('Monthly printing allowances: 200 black & white pages; 100 color pages. Overages: $0.10/page B&W; $0.20/page color.')
        add_bullet('Conference room overages beyond included hours: $50/hour, billed monthly.')
        add_bullet('Mail handling is included with all private office memberships.')

        # ── SECTION 4 ──────────────────────────────────────────────────────
        add_heading('4. Conduct & Responsibilities')
        add_bullet('Members shall conduct themselves in a professional, courteous, and cooperative manner at all times. Disruptive behavior, harassment, or conduct detrimental to the community may result in immediate termination of membership.')
        add_bullet('Noise: Use headphones for audio; avoid speakerphone calls or amplified sound in common areas. Keep voices at a conversational level.')
        add_bullet('Cleanliness: Wash all dishes immediately after use. Label personal food and beverages; unlabeled items will be discarded weekly. Clear personal items from common areas daily.')
        add_bullet('Safety: No hazardous materials, open flames, smoking, or pets anywhere on the premises. Report any safety hazards to management promptly.')
        add_bullet('Members are fully responsible for their own conduct and the conduct of any guests they bring onto the premises. Any damage caused by a member or their guest must be reported immediately and paid for in full.')
        add_bullet('Prohibited activities include, without limitation: pyramid or multi-level marketing schemes, harassment of any kind, unauthorized use of others\' personal or business information, theft, or display of inappropriate content.')

        # ── SECTION 5 ──────────────────────────────────────────────────────
        add_heading('5. Insurance & Liability')
        add_bullet('Qbix Centre does not provide insurance coverage for members\' personal property, equipment, or business assets. Members are strongly encouraged to obtain appropriate business and property insurance.')
        add_bullet('RoseAn Properties, LLC, its managers, officers, and affiliates shall not be liable for any theft, loss, damage, or injury to persons or property occurring on the premises, to the maximum extent permitted by law.')
        add_bullet('Member agrees to indemnify, defend, and hold harmless RoseAn Properties, LLC and Qbix Centre from and against any and all claims, damages, losses, or expenses (including reasonable attorneys\' fees) arising out of or related to Member\'s use of the premises or breach of this Agreement.')

        # ── SECTION 6 ──────────────────────────────────────────────────────
        add_heading('6. Termination')
        add_bullet('By Member: Written notice of termination must be provided at least thirty (30) days prior to the end of the then-current term. Membership fees remain due and payable through the end of the notice period. No pro-ration of fees for partial months.')
        add_bullet('By Management: Qbix Centre reserves the right to terminate any membership immediately and without refund for violation of this Agreement, the House Guidelines, or any conduct deemed detrimental to the community or the facility.')
        add_bullet('Upon termination, Member must return all keys and access devices, remove all personal property within 48 hours, and leave their office and any common areas in clean condition. Management may dispose of any property left after 48 hours.')

        # ── SECTION 7 ──────────────────────────────────────────────────────
        add_heading('7. General Provisions')
        add_bullet('Not a Lease: This Agreement grants a revocable license to use shared workspace and does not create a landlord-tenant relationship, leasehold interest, or any real property right of any kind.')
        add_bullet('Force Majeure: Services may be suspended without liability for events beyond management\'s reasonable control, including acts of God, utility failures, governmental orders, or civil unrest.')
        add_bullet('Modifications: Qbix Centre reserves the right to update the House Guidelines and membership policies at any time. Members will be notified of material changes by email.')
        add_bullet('Authority: The person signing this Agreement represents and warrants that they have full authority to bind themselves and/or their company to the terms herein.')
        add_bullet('Promotional Use: Member consents to Qbix Centre publishing their business name and general description in member directories and using non-identifiable photos of common areas for promotional purposes. Written consent is required for photos that identify individual members.')
        add_bullet('Dispute Resolution: The parties agree to attempt to resolve any dispute informally before pursuing legal remedies. This Agreement shall be governed by the laws of the State of Georgia.')
        add_bullet('Entire Agreement: This Agreement, together with the House Guidelines, constitutes the entire agreement between the parties and supersedes all prior negotiations, representations, or agreements.')

        # ── HOUSE GUIDELINES ──────────────────────────────────────────────
        add_heading('8. House Guidelines')
        add_body('The following guidelines are incorporated into and made part of this Agreement:')
        guidelines = [
            'Workstations & Equipment: Treat all furniture and equipment with care. Report damage promptly. Do not move furniture without management approval.',
            'Conference Room: Reserve in advance via qbixcentre.com. Clean up completely after each use — wipe the whiteboard, arrange chairs, and dispose of any trash.',
            'Kitchenette: Label all food and beverages with your name and date. Unclaimed items will be discarded each Friday. Wash dishes immediately — do not leave in the sink.',
            'Noise & Privacy: Keep voices low in common areas. Use headphones for all audio. Conference rooms are available for calls and meetings.',
            'Visitors & Clients: You are welcome to receive clients. Guide them directly to your office or a conference room. Do not leave guests unattended in common areas.',
            'Internet: Use the network responsibly and lawfully. Do not conduct illegal downloads or activities. Bandwidth-intensive personal streaming is not permitted.',
            'Printing: Print thoughtfully. Overages will be added to your next invoice.',
            'Cleanliness: Clear your desk daily. Use recycling bins for paper. Use the secure shredding bin for confidential documents.',
            'Smoke-Free Facility: No smoking, vaping, or e-cigarettes anywhere on the property, including the parking lot.',
            'Fragrances: Please be considerate of shared air — avoid strong perfumes, candles, or incense.',
            'Comfort: Thermostat adjustments should be discussed with management. Do not make changes without approval.',
        ]
        for g in guidelines:
            add_bullet(g)

        # ── SIGNATURE PAGE ────────────────────────────────────────────────
        doc.add_page_break()

        sig_title = doc.add_paragraph()
        sig_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sig_title.paragraph_format.space_after = Pt(4)
        str_ = sig_title.add_run('SIGNATURE PAGE')
        str_.bold = True
        str_.font.size = Pt(13)
        str_.font.color.rgb = RGBColor(0x1a, 0x27, 0x44)

        sig_sub = doc.add_paragraph()
        sig_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sig_sub.paragraph_format.space_after = Pt(16)
        ssr = sig_sub.add_run('Qbix Centre Membership Agreement')
        ssr.font.size = Pt(10)
        ssr.italic = True

        add_heading('Member Acknowledgment')
        add_body(
            'By signing below, Member acknowledges that they have read, understand, and agree to all '
            'terms and conditions of this Membership Agreement and House Guidelines, and that they '
            'have the authority to enter into this Agreement on behalf of themselves and/or their company.'
        )
        doc.add_paragraph()

        add_field_row('Member / Company', member.get('name', ''))
        doc.add_paragraph()
        add_sig_line('Member Signature')
        add_sig_line('Print Name', member.get('name', ''))
        add_sig_line('Title / Position')
        add_sig_line('Date')
        doc.add_paragraph()

        add_heading('Qbix Centre Acceptance')
        doc.add_paragraph()
        add_sig_line('Authorized Signature')
        add_sig_line('Print Name')
        add_sig_line('Title', 'Manager, RoseAn Properties, LLC')
        add_sig_line('Date')
        doc.add_paragraph()

        add_field_row('Initial Term Start Date', start_str)
        add_field_row('Office(s) Assigned', office_str)
        add_field_row('Monthly Dues', dues_str)
        add_field_row('Deposit Collected', deposit_str)

        # ── BACKGROUND CHECK AUTHORIZATION ────────────────────────────────
        doc.add_page_break()

        bc_title = doc.add_paragraph()
        bc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bc_title.paragraph_format.space_after = Pt(4)
        bcr = bc_title.add_run('CONFIDENTIAL BACKGROUND CHECK AUTHORIZATION')
        bcr.bold = True
        bcr.font.size = Pt(13)
        bcr.font.color.rgb = RGBColor(0x1a, 0x27, 0x44)

        bc_fee = doc.add_paragraph()
        bc_fee.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bc_fee.paragraph_format.space_after = Pt(16)
        bfr = bc_fee.add_run('Non-refundable fee: $35 per cardholder  |  Required prior to access')
        bfr.font.size = Pt(10)
        bfr.italic = True

        add_heading('Applicant Information')
        add_sig_line('Print Full Legal Name', member.get('name', ''))
        add_sig_line('Former Name(s) & Dates Used')
        add_sig_line('Social Security Number')
        add_sig_line('Date of Birth')
        add_sig_line('Current Address (include move-in month/year)')
        add_sig_line('Previous Address #1')
        add_sig_line('Previous Address #2')
        doc.add_paragraph()

        add_heading('Identification Required')
        add_body('A legible copy of your current driver\'s license or government-issued photo ID must be attached to this authorization. Please attach a copy below or submit separately.')
        doc.add_paragraph()

        add_heading('Authorization Statement')
        add_body(
            'I hereby authorize RoseAn Properties, LLC and its designated agents and representatives '
            'to conduct a comprehensive background investigation, including but not limited to: '
            'verification of credit history, residential history, employment history, educational '
            'background, professional references, criminal and civil court records, driving records, '
            'and any other public records deemed relevant. I understand that this authorization '
            'remains in effect until revoked in writing, and that a photocopy of this authorization '
            'shall be as valid as the original.'
        )
        doc.add_paragraph()
        add_sig_line('Applicant Signature')
        add_sig_line('Date')

        # ── AUTO-DRAFT AUTHORIZATION ───────────────────────────────────────
        doc.add_page_break()

        ad_title = doc.add_paragraph()
        ad_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ad_title.paragraph_format.space_after = Pt(4)
        adr = ad_title.add_run('AUTO-DRAFT AUTHORIZATION')
        adr.bold = True
        adr.font.size = Pt(13)
        adr.font.color.rgb = RGBColor(0x1a, 0x27, 0x44)

        ad_sub = doc.add_paragraph()
        ad_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ad_sub.paragraph_format.space_after = Pt(16)
        adsr = ad_sub.add_run('Bill.com, Inc. on behalf of RoseAn Properties, LLC  |  Required for all memberships')
        adsr.font.size = Pt(10)
        adsr.italic = True

        add_heading('Banking Information')
        add_sig_line('Account Holder Name', member.get('name', ''))
        add_sig_line('Bank Name')
        add_sig_line('Account Number')
        add_sig_line('Routing Number')
        add_body('(Attach a voided check if preferred in lieu of completing account information above.)')
        doc.add_paragraph()

        add_heading('Monthly Draft Amount')
        add_field_row('Authorized Monthly Amount', dues_str, bold_value=True)
        add_body('Draft date: 1st of each month. If the 1st falls on a weekend or holiday, the draft will process on the next business day.')
        doc.add_paragraph()

        add_heading('Authorization')
        add_body(
            f'I/We authorize Bill.com, Inc., on behalf of RoseAn Properties, LLC (Qbix Centre), '
            f'to initiate recurring ACH debit entries to the bank account identified above in the '
            f'amount of {dues_str}, beginning {start_str}. This authorization will remain in full '
            f'force and effect until I/we notify Qbix Centre in writing at least ten (10) business '
            f'days prior to the desired cancellation date. I/We have the right to stop payment of '
            f'any debit entry by notifying my/our bank at least three (3) business days prior to '
            f'the scheduled debit date.'
        )
        doc.add_paragraph()
        add_sig_line('Account Holder Signature')
        add_sig_line('Print Name', member.get('name', ''))
        add_sig_line('Date')

        # Save
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        safe_name = ''.join(c if c.isalnum() or c in ' _-' else '_' for c in member['name'])
        filename  = f"Qbix_Agreement_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx"

        # Mark agreement as generated
        member['agreementSent'] = datetime.now().strftime('%Y-%m-%d')
        save_data(data)

        return send_file(buf, as_attachment=True,
                         download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except ImportError:
        return jsonify({'ok': False, 'error': 'python-docx not installed on server'}), 500
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

# ── Newsletter generator (AI-assisted) ───────────────────────────────────────

@app.route('/admin/api/generate-newsletter', methods=['POST'])
@login_required
def generate_newsletter():
    if not ANTHROPIC_API_KEY:
        return jsonify({'ok': False, 'error': 'Anthropic API key not configured'}), 400

    data   = get_db()
    active = [m for m in data['members'] if m['status'] == 'Active']
    occ    = len([o for o in data['offices'] if o['status'] == 'Occupied'])
    vac    = len([o for o in data['offices'] if o['status'] == 'Vacant'])
    month  = datetime.now().strftime('%B %Y')

    context = (
        f"Qbix Centre is a professional coworking space in Macon, Georgia at 500A Northside Crossing. "
        f"Current stats: {occ} offices occupied, {vac} vacant, {len(active)} active members. "
        f"Month: {month}. "
        f"Amenities: 24/7 access, AT&T Fiber, Starbucks coffee, furnished offices, conference room, free parking."
    )

    custom_notes = request.json.get('notes', '')

    try:
        import urllib.request
        import json as json_mod

        payload = {
            'model': 'claude-sonnet-4-20250514',
            'max_tokens': 1000,
            'messages': [{
                'role': 'user',
                'content': (
                    f'Write a warm, friendly, professional monthly newsletter for Qbix Centre. '
                    f'Keep it concise — 3-4 short paragraphs. Include a welcoming opener, '
                    f'a community update, any relevant seasonal note, and a friendly closing. '
                    f'Context: {context}. '
                    f'Additional notes from manager: {custom_notes}. '
                    f'Format as HTML suitable for email. Use a warm, community-focused tone.'
                )
            }]
        }

        req = urllib.request.Request(
            'https://api.anthropic.com/v1/messages',
            data=json_mod.dumps(payload).encode(),
            headers={
                'Content-Type': 'application/json',
                'x-api-key': ANTHROPIC_API_KEY,
                'anthropic-version': '2023-06-01',
            }
        )
        with urllib.request.urlopen(req, timeout=30) as resp:
            result = json_mod.loads(resp.read())
            draft = result['content'][0]['text']
            return jsonify({'ok': True, 'draft': draft})

    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/admin/api/publish-newsletter', methods=['POST'])
@login_required
def publish_newsletter():
    """Save newsletter and optionally email to all active members."""
    data    = get_db()
    subject = request.json.get('subject', f'Qbix Centre Newsletter — {datetime.now().strftime("%B %Y")}')
    body    = request.json.get('body', '')
    send    = request.json.get('send', False)

    post_id = '_' + secrets.token_hex(4)
    post = {
        'id':      post_id,
        'subject': subject,
        'body':    body,
        'date':    datetime.now().isoformat(),
        'sent':    send,
    }
    data.setdefault('newsletter', []).append(post)
    save_data(data)

    if send:
        active_with_email = [m for m in data['members']
                             if m['status'] == 'Active' and m.get('email')]
        sent_count = 0
        for m in active_with_email:
            ok = send_email(m['email'], m['name'], subject, body)
            if ok:
                sent_count += 1
        return jsonify({'ok': True, 'sent': sent_count, 'postId': post_id})

    return jsonify({'ok': True, 'sent': 0, 'postId': post_id})

# ── Booking management (admin view) ──────────────────────────────────────────

@app.route('/admin/bookings')
@login_required
def admin_bookings():
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/api/booking-cancel', methods=['POST'])
@login_required
def admin_cancel_booking():
    booking_id = request.json.get('bookingId', '')
    data = get_db()
    booking = next((b for b in data.get('bookings', []) if b['id'] == booking_id), None)
    if not booking:
        return jsonify({'ok': False}), 404
    booking['status'] = 'cancelled'
    save_data(data)
    return jsonify({'ok': True})

# ── Monthly usage email (called by scheduler or manually) ────────────────────

@app.route('/admin/api/send-monthly-usage', methods=['POST'])
@login_required
def send_monthly_usage():
    data  = get_db()
    month = request.json.get('month', datetime.now().month)
    year  = request.json.get('year',  datetime.now().year)

    active_with_email = [m for m in data['members']
                         if m['status'] == 'Active' and m.get('email')]
    sent = 0
    month_name = datetime(year, month, 1).strftime('%B %Y')

    for member in active_with_email:
        # Count their bookings that month
        member_bookings = [
            b for b in data.get('bookings', [])
            if b.get('memberEmail','').lower() == member.get('email','').lower()
            and b.get('year') == year
            and b.get('month') == month
            and b.get('status') != 'cancelled'
        ]

        # Calculate hours used
        total_minutes = 0
        for b in member_bookings:
            try:
                s = datetime.strptime(b['start'], '%H:%M')
                e = datetime.strptime(b['end'],   '%H:%M')
                total_minutes += int((e - s).total_seconds() / 60)
            except Exception:
                pass
        hours_used     = round(total_minutes / 60, 1)
        included       = hours_included(data, member['name'])
        hours_remaining = max(0, included - hours_used)

        if hours_used == 0 and not member_bookings:
            # Skip members who didn't use the room
            continue

        booking_list = ''.join(
            f'<li>{b["date"]} {b["start"]}–{b["end"]}: {b.get("title","Meeting")}</li>'
            for b in member_bookings
        )

        send_email(
            member['email'], member['name'],
            f'Your Qbix Centre Conference Room Summary — {month_name}',
            f'<p>Hi {member["name"]},</p>'
            f'<p>Here\'s a cheerful summary of your conference room use in {month_name}! 🎉</p>'
            f'<ul>{booking_list}</ul>'
            f'<p><b>Total hours used:</b> {hours_used} of your {included} included hours</p>'
            f'<p><b>Hours remaining this month:</b> {hours_remaining}</p>'
            f'<p>Thank you for being part of the Qbix Centre community. See you next month!</p>'
            f'<p style="color:#666;font-size:12px">Questions? Reply to this email or call (478) 216-2876</p>',
        )
        sent += 1

    return jsonify({'ok': True, 'sent': sent})


# ── Setup route (first run only) ─────────────────────────────────────────────

@app.route('/admin/setup', methods=['GET', 'POST'])
def admin_setup():
    """First-run setup to set admin password."""
    if os.environ.get('ADMIN_PASSWORD_HASH'):
        return redirect(url_for('admin_login'))

    if request.method == 'POST':
        password = request.form.get('password', '')
        confirm  = request.form.get('confirm', '')
        if password != confirm:
            flash('Passwords do not match.', 'error')
        elif len(password) < 8:
            flash('Password must be at least 8 characters.', 'error')
        else:
            hashed = hash_password(password)
            flash(
                f'Setup complete! Add this to your Railway environment variables: '
                f'ADMIN_PASSWORD_HASH={hashed}',
                'success'
            )
    return render_template('admin/setup.html')


@app.context_processor
def inject_now():
    return {'now': datetime.now()}


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8765))
    app.run(host='0.0.0.0', port=port, debug=False)
